from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
from datetime import datetime
import os

def generate_pdf_report(data, report_type, output_path):
    """
    Generate PDF report with the given data
    :param data: Dictionary containing report data
    :param report_type: Type of report (attendance, grades, etc.)
    :param output_path: Path where the PDF should be saved
    """
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Add title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30
    )
    title = Paragraph(f"{report_type.title()} Report", title_style)
    story.append(title)
    story.append(Spacer(1, 12))

    # Add date
    date_style = ParagraphStyle(
        'DateStyle',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=20
    )
    date = Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", date_style)
    story.append(date)
    story.append(Spacer(1, 12))

    # Add content based on report type
    if report_type == 'attendance':
        _add_attendance_content(story, data, styles)
    elif report_type == 'grades':
        _add_grades_content(story, data, styles)
    # Add more report types as needed

    doc.build(story)

def generate_docx_report(data, report_type, output_path):
    """
    Generate DOCX report with the given data
    :param data: Dictionary containing report data
    :param report_type: Type of report (attendance, grades, etc.)
    :param output_path: Path where the DOCX should be saved
    """
    doc = Document()
    
    # Add title
    doc.add_heading(f"{report_type.title()} Report", 0)
    
    # Add date
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()  # Add space

    # Add content based on report type
    if report_type == 'attendance':
        _add_attendance_content_docx(doc, data)
    elif report_type == 'grades':
        _add_grades_content_docx(doc, data)
    # Add more report types as needed

    doc.save(output_path)

def _add_attendance_content(story, data, styles):
    """Helper function to add attendance data to PDF"""
    if 'class_info' in data:
        # Add class info
        class_info = Paragraph(f"Class: {data['class_info']['name']}", styles['Heading2'])
        story.append(class_info)
        story.append(Spacer(1, 12))

    if 'students' in data:
        # Create table data
        table_data = [['Student ID', 'Name', 'Status', 'Date']]
        for student in data['students']:
            table_data.append([
                student.get('id', ''),
                student.get('name', ''),
                student.get('status', ''),
                student.get('date', '')
            ])

        # Create and style table
        table = Table(table_data, colWidths=[1*inch, 2*inch, 1*inch, 1.5*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(table)

def _add_grades_content(story, data, styles):
    """Helper function to add grades data to PDF"""
    if 'exam_info' in data:
        # Add exam info
        exam_info = Paragraph(f"Exam: {data['exam_info']['name']}", styles['Heading2'])
        story.append(exam_info)
        story.append(Spacer(1, 12))

    if 'students' in data:
        # Create table data
        table_data = [['Student ID', 'Name', 'Grade', 'Marks']]
        for student in data['students']:
            table_data.append([
                student.get('id', ''),
                student.get('name', ''),
                student.get('grade', ''),
                str(student.get('marks', ''))
            ])

        # Create and style table
        table = Table(table_data, colWidths=[1*inch, 2*inch, 1*inch, 1*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(table)

def _add_attendance_content_docx(doc, data):
    """Helper function to add attendance data to DOCX"""
    if 'class_info' in data:
        doc.add_heading(f"Class: {data['class_info']['name']}", level=1)

    if 'students' in data:
        # Add table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Student ID'
        header_cells[1].text = 'Name'
        header_cells[2].text = 'Status'
        header_cells[3].text = 'Date'

        # Add data rows
        for student in data['students']:
            row_cells = table.add_row().cells
            row_cells[0].text = student.get('id', '')
            row_cells[1].text = student.get('name', '')
            row_cells[2].text = student.get('status', '')
            row_cells[3].text = student.get('date', '')

def _add_grades_content_docx(doc, data):
    """Helper function to add grades data to DOCX"""
    if 'exam_info' in data:
        doc.add_heading(f"Exam: {data['exam_info']['name']}", level=1)

    if 'students' in data:
        # Add table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Student ID'
        header_cells[1].text = 'Name'
        header_cells[2].text = 'Grade'
        header_cells[3].text = 'Marks'

        # Add data rows
        for student in data['students']:
            row_cells = table.add_row().cells
            row_cells[0].text = student.get('id', '')
            row_cells[1].text = student.get('name', '')
            row_cells[2].text = student.get('grade', '')
            row_cells[3].text = str(student.get('marks', ''))
